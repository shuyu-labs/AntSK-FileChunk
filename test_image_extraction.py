#!/usr/bin/env python3
"""
测试图片提取功能的脚本
"""

import os
import sys
import logging
from pathlib import Path

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent / "src"))

from antsk_filechunk.unified_document_parser import UnifiedDocumentParser
from antsk_filechunk.enhanced_semantic_chunker import EnhancedSemanticChunker

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def test_image_extraction():
    """测试图片提取功能"""
    
    # 创建解析器
    parser = UnifiedDocumentParser(
        image_base_url="http://localhost:8000",
        image_save_dir="./static/images"
    )
    
    # 创建切片器
    chunker = EnhancedSemanticChunker()
    
    # 测试文件路径（需要用户提供包含图片的文档）
    test_files = [
        # 用户可以在这里添加测试文件路径
        # "path/to/test_document_with_images.docx",
        # "path/to/test_document_with_images.pdf"
    ]
    
    if not test_files:
        logger.info("请在test_files列表中添加包含图片的测试文档路径")
        return
    
    for file_path in test_files:
        if not os.path.exists(file_path):
            logger.warning(f"测试文件不存在: {file_path}")
            continue
            
        try:
            logger.info(f"正在测试文件: {file_path}")
            
            # 1. 解析文档
            document_content = parser.parse_file(file_path)
            
            logger.info(f"解析结果:")
            logger.info(f"  - 段落数: {len(document_content.paragraphs)}")
            logger.info(f"  - 表格数: {len(document_content.tables)}")
            logger.info(f"  - 图片数: {len(document_content.images)}")
            
            # 打印图片信息
            for i, image in enumerate(document_content.images):
                logger.info(f"  图片 {i+1}: {image.get('filename')} -> {image.get('url')}")
            
            # 2. 进行语义切片
            chunks = chunker.process_document(document_content)
            
            logger.info(f"切片结果:")
            logger.info(f"  - 切片总数: {len(chunks)}")
            
            # 检查哪些切片包含图片
            image_chunks = []
            for i, chunk in enumerate(chunks):
                if chunk.chunk_type in ['image_content', 'mixed_content'] or '![' in chunk.content:
                    image_chunks.append(i)
                    logger.info(f"  切片 {i+1} 包含图片内容 (类型: {chunk.chunk_type})")
                    
                    # 显示切片内容的前100个字符
                    content_preview = chunk.content[:100] + "..." if len(chunk.content) > 100 else chunk.content
                    logger.info(f"    内容预览: {content_preview}")
            
            if image_chunks:
                logger.info(f"成功！共有 {len(image_chunks)} 个切片包含图片内容")
            else:
                logger.warning("警告：没有切片包含图片内容，可能存在问题")
            
            # 3. 检查markdown内容
            if '![' in document_content.markdown_content:
                logger.info("✓ 文档的markdown内容包含图片引用")
            else:
                logger.warning("✗ 文档的markdown内容不包含图片引用")
                
        except Exception as e:
            logger.error(f"测试文件 {file_path} 时出错: {e}")
            import traceback
            traceback.print_exc()

def create_test_document():
    """创建一个包含图片的测试文档（示例）"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        # 创建一个简单的测试文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('测试文档 - 包含图片', 0)
        
        # 添加段落
        doc.add_paragraph('这是第一个段落，下面会插入一张图片。')
        
        # 注意：这里需要一个实际的图片文件
        # doc.add_picture('path/to/test_image.png', width=Inches(2))
        
        doc.add_paragraph('这是图片后的段落。')
        
        # 保存文档
        test_doc_path = "test_document_with_images.docx"
        doc.save(test_doc_path)
        
        logger.info(f"创建测试文档: {test_doc_path}")
        logger.info("注意：请手动在文档中插入图片后再进行测试")
        
    except ImportError:
        logger.warning("需要安装python-docx库来创建测试文档")
    except Exception as e:
        logger.error(f"创建测试文档失败: {e}")

if __name__ == "__main__":
    print("图片提取功能测试脚本")
    print("=" * 50)
    
    # 检查是否有测试文件
    if len(sys.argv) > 1:
        test_files = sys.argv[1:]
        logger.info(f"使用命令行提供的测试文件: {test_files}")
        
        # 创建解析器和切片器
        parser = UnifiedDocumentParser()
        chunker = EnhancedSemanticChunker()
        
        for file_path in test_files:
            if os.path.exists(file_path):
                try:
                    logger.info(f"\n正在测试: {file_path}")
                    
                    # 解析文档
                    document_content = parser.parse_file(file_path)
                    
                    # 显示解析结果
                    print(f"解析结果:")
                    print(f"  段落数: {len(document_content.paragraphs)}")
                    print(f"  表格数: {len(document_content.tables)}")
                    print(f"  图片数: {len(document_content.images)}")
                    
                    if document_content.images:
                        print("  图片列表:")
                        for i, img in enumerate(document_content.images):
                            print(f"    {i+1}. {img.get('filename')} ({img.get('format')}) -> {img.get('url')}")
                    
                    # 进行切片
                    chunks = chunker.process_document(document_content)
                    print(f"  切片数: {len(chunks)}")
                    
                    # 检查包含图片的切片
                    image_chunks = 0
                    for chunk in chunks:
                        if '![' in chunk.content:
                            image_chunks += 1
                    
                    print(f"  包含图片的切片数: {image_chunks}")
                    
                    if image_chunks > 0:
                        print("✓ 图片提取和切片功能正常工作")
                    else:
                        print("⚠ 没有检测到包含图片的切片")
                        
                except Exception as e:
                    print(f"✗ 处理文件时出错: {e}")
            else:
                print(f"✗ 文件不存在: {file_path}")
    else:
        print("使用方法:")
        print("  python test_image_extraction.py <文档文件路径1> [文档文件路径2] ...")
        print("\n示例:")
        print("  python test_image_extraction.py test.docx test.pdf")
        print("\n或者修改脚本中的test_files列表来指定测试文件")
